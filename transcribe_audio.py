"""

~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
Audio Transcription Pipeline for Echelon NOS
~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

Authors: Christian Bader
October 2024 - Present

This script iterates through audio/video files in a Google Drive folder and transcribes them. 
Audio files are saved as .mp3 files in a new folder, Transcripts are saved as .docx in a new folder.
.docx files are assigned metadata properties like transcription date/time, seconds transcribed, etc.
Custom properties set on a file in Google Drive using the API are not visible through the 
Google Drive web interface. To access these properties, you need to use the Google Drive API.

Deployed on Streamlit Cloud.

"""

# Echelon imports
# e.g. from utilities import function
#from gdrive_functions import gd_list_audio_video_files,gd_download_file,gd_upload_file,gd_get_file_properties,gd_update_file_properties,gd_move_file_between_folders,gd_get_shareable_link

# Standard python library imports
import os
import io
import re
import json
import shutil
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import base64

# Open source imports
import streamlit as st
from docx import Document
from pydub import AudioSegment
from moviepy.editor import VideoFileClip

# API imports
from openai import OpenAI
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request

# Define OpenAI scopes/credentials, initialize client
os.environ['OPENAI_API_KEY'] = st.secrets["openai_api_key"]
client = OpenAI()

# Define Google scopes/credentials, initialize client
SCOPES = [
    'https://www.googleapis.com/auth/drive',
    'https://www.googleapis.com/auth/documents',
    'https://www.googleapis.com/auth/spreadsheets'
]
gcp_secrets = st.secrets["gcp_service_account"]
creds = service_account.Credentials.from_service_account_info(
    gcp_secrets,
    scopes=SCOPES
)
drive_service = build('drive', 'v3', credentials=creds)
docs_service = build('docs', 'v1', credentials=creds)
sheets_service = build('sheets', 'v4', credentials=creds)

def get_gmail_service():
    creds = Credentials(
        None,
        refresh_token=st.secrets["gmail"]["refresh_token"],
        token_uri="https://oauth2.googleapis.com/token",
        client_id=st.secrets["gmail"]["client_id"],
        client_secret=st.secrets["gmail"]["client_secret"],
        scopes=["https://www.googleapis.com/auth/gmail.send"]
    )
    creds.refresh(Request())
    service = build('gmail', 'v1', credentials=creds)
    return service

# Define Google Drive folder and spreadsheet IDs
#PRODUCTION IDs
# GD_FOLDER_ID_UNPROCESSED_AUDIO = st.secrets["GD_FOLDER_ID_UNPROCESSED_AUDIO_PROD"]
# GD_FOLDER_ID_TRANSCRIBED_AUDIO = st.secrets["GD_FOLDER_ID_TRANSCRIBED_AUDIO_PROD"]
# GD_FOLDER_ID_TRANSCRIBED_TEXT = st.secrets["GD_FOLDER_ID_TRANSCRIBED_TEXT_PROD"]
# GD_FOLDER_ID_PROCESSED_RAW_AUDIO = st.secrets["GD_FOLDER_ID_PROCESSED_RAW_AUDIO_PROD"]
# GD_SPREADSHEET_ID_INGRESS_LOG = st.secrets["GD_SPREADSHEET_ID_INGRESS_LOG_PROD"]
# GD_SHEET_NAME_INGRESS_LOG = 'transcribe_audio'

#TEST IDs
GD_FOLDER_ID_UNPROCESSED_AUDIO = st.secrets["GD_FOLDER_ID_UNPROCESSED_AUDIO_TEST"]
GD_FOLDER_ID_TRANSCRIBED_AUDIO = st.secrets["GD_FOLDER_ID_TRANSCRIBED_AUDIO_TEST"]
GD_FOLDER_ID_TRANSCRIBED_TEXT = st.secrets["GD_FOLDER_ID_TRANSCRIBED_TEXT_TEST"]
GD_FOLDER_ID_PROCESSED_RAW_AUDIO = st.secrets["GD_FOLDER_ID_PROCESSED_RAW_AUDIO_TEST"]
GD_SPREADSHEET_ID_INGRESS_LOG = st.secrets["GD_SPREADSHEET_ID_INGRESS_LOG_TEST"]
GD_SHEET_NAME_INGRESS_LOG = 'transcribe_audio'

# Define functions that interact with local repo

def convert_to_mp3(input_file, mime_type):
    """
    Converts any audio or video file to MP3 format using MIME type for identification and returns the file path
    along with the audio duration in seconds.

    Parameters:
        input_file (str): The path to the input audio or video file.
        mime_type (str): The MIME type of the input file.

    Returns:
        tuple: The path to the converted .mp3 file and the duration in seconds.

    Raises:
        Exception: If there is an error during the conversion process.
    """
    import os
    import shutil
    from pydub import AudioSegment

    # Determine the output file path by replacing the extension with .mp3
    base, _ = os.path.splitext(input_file)
    output_file = base + '.mp3'

    try:
        if mime_type == 'audio/mpeg':
            # If the file is already an MP3
            if os.path.abspath(input_file) != os.path.abspath(output_file):
                # Copy the file to the output path if it's not the same file
                shutil.copy(input_file, output_file)
            else:
                # Input and output files are the same; no action needed
                pass
        elif mime_type.startswith('video/') or mime_type.startswith('audio/'):
            # Use FFmpeg to extract audio from video or convert audio to MP3
            extract_audio_with_ffmpeg(input_file, output_file)
        else:
            # Unsupported file type
            print(f"Unsupported file type: {mime_type}")
            return None, None

        # Load the converted MP3 file with pydub to get the duration
        audio = AudioSegment.from_mp3(output_file)
        duration_seconds = len(audio) / 1000  # pydub returns duration in milliseconds
    except Exception as e:
        print(f"Error converting file {input_file}: {e}")
        raise

    # Return the output file path and duration in seconds
    return output_file, duration_seconds


def extract_audio_with_ffmpeg(input_file, output_file):
    """
    Extracts audio from a video file or converts an audio file to MP3 using FFmpeg.

    Parameters:
        input_file (str): The path to the input file.
        output_file (str): The path to save the output MP3 file.

    Raises:
        subprocess.CalledProcessError: If the FFmpeg command fails.
    """
    import subprocess

    # Construct the FFmpeg command to extract audio and convert to MP3
    command = [
        'ffmpeg',
        '-i', input_file,
        '-vn',                 # Disable video recording (process audio only)
        '-acodec', 'libmp3lame',  # Use the MP3 audio codec
        '-q:a', '2',              # Set audio quality (2 is high quality)
        output_file
    ]

    try:
        # Run the FFmpeg command
        subprocess.run(command, check=True)
    except subprocess.CalledProcessError as e:
        print(f"FFmpeg failed with error: {e}")
        raise


def rename_file(input_file_path, new_file_name):
    """
    Renames a file to a new name in the same directory.

    Parameters:
        input_file_path (str): The full path to the original file.
        new_file_name (str): The new file name (with extension).

    Returns:
        str: The full path to the renamed file.
    """
    dir_name = os.path.dirname(input_file_path)
    new_file_path = os.path.join(dir_name, new_file_name)
    os.rename(input_file_path, new_file_path)

    return new_file_path


# Define functions that interact with Google Docs + Drive

def gd_get_file_creation_date(file_id):
    """
    Retrieves the original upload date (creation date) of a file from Google Drive.

    Parameters:
        file_id (str): The ID of the file.

    Returns:
        str: The creation date and time in ISO 8601 format (e.g., '2023-11-04T12:34:56.789Z').

    Raises:
        Exception: If there is an error retrieving the creation date.
    """
    try:
        file = drive_service.files().get(fileId=file_id, fields='createdTime').execute()
        created_time = file.get('createdTime')
        return created_time
    except Exception as e:
        print(f"Error retrieving creation date for file {file_id}: {str(e)}")
        return None

def gd_list_audio_video_files(folder_id):
    """
    Lists all audio and video files in a Google Drive folder.

    Parameters:
        folder_id (str): The ID of the Google Drive folder.

    Returns:
        list: A list of files with their 'id', 'name', and 'mimeType'.
    """
    query = f"'{folder_id}' in parents and (mimeType contains 'audio/' or mimeType contains 'video/')"
    results = drive_service.files().list(q=query, fields="files(id, name, mimeType)").execute()
    files = results.get('files', [])
    return files


def gd_download_file(file_id, file_name):
    """
    Downloads a file from Google Drive.

    Parameters:
        file_id (str): The ID of the file to download.
        file_name (str): The name to save the file as locally.

    Returns:
        str: The local path to the downloaded file.
    """
    request = drive_service.files().get_media(fileId=file_id)
    fh = io.FileIO(file_name, 'wb')
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()
        print(f"Download {file_name}: {int(status.progress() * 100)}%.")
    return file_name


def gd_upload_file(file_path, folder_id, mime_type):
    """
    Uploads a file to a specified Google Drive folder.

    Parameters:
        file_path (str): The local path to the file to upload.
        folder_id (str): The ID of the destination Google Drive folder.
        mime_type (str): The MIME type of the file.

    Returns:
        str: The ID of the uploaded file in Google Drive.
    """
    file_metadata = {
        'name': os.path.basename(file_path),
        'parents': [folder_id]
    }

    media = MediaFileUpload(file_path, mimetype=mime_type)
    uploaded_file = drive_service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id'
    ).execute()
    return uploaded_file.get('id')


def gd_get_file_properties(file_id):
    """
    Retrieves the properties of a file from Google Drive.

    Parameters:
        file_id (str): The ID of the file.

    Returns:
        dict: A dictionary containing the file's properties.
    """
    file = drive_service.files().get(fileId=file_id, fields='properties').execute()
    properties = file.get('properties', {})
    return properties


def gd_update_file_properties(file_id, new_properties):
    """
    Updates the properties of a file in Google Drive.

    Parameters:
        file_id (str): The ID of the file.
        new_properties (dict): A dictionary of new properties to set.

    Returns:
        dict: The updated file resource.
    """
    file_metadata = {
        'properties': new_properties
    }
    updated_file = drive_service.files().update(
        fileId=file_id,
        body=file_metadata,
        fields='id, properties'
    ).execute()
    return updated_file

def gd_move_file_between_folders(file_id, target_folder_id):
    """
    Moves a file to a different Google Drive folder.

    Parameters:
        file_id (str): The ID of the file to move.
        target_folder_id (str): The ID of the destination folder.

    Returns:
        None
    """
    try:
        # Retrieve the existing parents to remove
        file = drive_service.files().get(fileId=file_id, fields='parents').execute()
        previous_parents = ",".join(file.get('parents'))

        # Move the file to the new folder
        drive_service.files().update(
            fileId=file_id,
            addParents=target_folder_id,
            removeParents=previous_parents,
            fields='id, parents'
        ).execute()

        print(f"File ID {file_id} moved to folder ID {target_folder_id}")
    except Exception as e:
        print(f"Error moving file {file_id}: {str(e)}")


def gd_get_shareable_link(file_id):
    """
    Creates a shareable link for a Google Drive file.

    Parameters:
        file_id (str): The ID of the file.

    Returns:
        str: The shareable link to the file.
    """
    try:
        # Update file permissions to make it shareable
        permission = {
            'type': 'anyone',
            'role': 'reader'
        }
        drive_service.permissions().create(fileId=file_id, body=permission).execute()

        # Get the shareable link
        file = drive_service.files().get(fileId=file_id, fields='webViewLink').execute()
        return file.get('webViewLink')
    except Exception as e:
        print(f"Error getting shareable link for file {file_id}: {str(e)}")
        return None

# Define functions that leverage OpenAI API

def transcribe(audio_file_path):
    """
    Transcribes an audio file to text using OpenAI Whisper.

    Parameters:
        audio_file_path (str): The local path to the audio file.

    Returns:
        str: The transcribed text.
    """
    with open(audio_file_path, 'rb') as audio_file:
        transcription = client.audio.transcriptions.create(
            model="whisper-1", 
            file=audio_file,
        )
    return transcription.text


def openai_llm_call(system_prompt, user_prompt, client):
    """
    Formats the transcription using the GPT-4 API.

    Parameters:
        system_prompt (str): The system prompt to guide the model's behavior.
        user_prompt (str): The user's message content, typically the raw transcription.
        client: The OpenAI API client instance.

    Returns:
        str: The formatted transcription returned by the GPT model.
    """
    try:
        # Send the prompts to GPT-4 for formatting
        completion = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ]
        )
        # Extract the formatted transcription from the response
        output = completion.choices[0].message.content if completion.choices[0].message else ""
    except Exception as e:
        # Handle any exceptions that occur during the GPT-4 API call
        output = "LLM Processing Failed. Use ChatGPT manually."
    return output


# Streamlit UI
st.set_page_config(
        page_title="NOS Speech2Text",
        page_icon="Echelon_Icon_Sky Blue.png",
)
st.image("Echelon_Icon_Sky Blue.png", caption="The Home for Aliens", width = 125)
st.title("NOS - Transcribe Audio Files")
st.write("Custom Built for Kerri Faber")
st.write("Once you have uploaded your files to the folder linked below, click the 'Transcribe Audio Files' button to transcribe. Full instructions are available on Notion.")
upload_folder_link = gd_get_shareable_link(GD_FOLDER_ID_UNPROCESSED_AUDIO)
st.markdown(f'[Upload Folder]({upload_folder_link})')
st.markdown('[Notion](https://www.notion.so/Pulse-4799295f90594380b55f75e0d78dbb03?p=11b9668a26d680e39d57e8243d8f7178&pm=s)')

# Add a reset button
if st.button('Reset App'):
    st.query_params.clear()  # Simulate a reset by clearing query parameters

if st.button('Transcribe Audio Files'):
    st.write("Transcription started...")

    try:
        gd_audio_files = gd_list_audio_video_files(GD_FOLDER_ID_UNPROCESSED_AUDIO)
        processed_files_count = 0
        gd_file_count = len(gd_audio_files)
        st.write(f"Found {gd_file_count} audio files to transcribe.")

        for file in gd_audio_files:
            gd_input_audio_file_id = file['id']
            gd_input_audio_file_name = file['name']  # Original file name
            gd_input_audio_file_mimeType = file['mimeType']
            gd_input_audio_file_link = gd_get_shareable_link(gd_input_audio_file_id)

            # Get the original upload date
            gd_input_audio_file_createdTime = gd_get_file_creation_date(gd_input_audio_file_id)

            # Convert to formatted date/time string
            if gd_input_audio_file_createdTime:
                gd_input_audio_file_created_datetime = datetime.strptime(
                    gd_input_audio_file_createdTime, '%Y-%m-%dT%H:%M:%S.%fZ'
                )
                datetime_uploaded = gd_input_audio_file_created_datetime.strftime('%Y-%m-%d %H:%M:%S')
            else:
                datetime_uploaded = 'Unknown'
            
            datetime_uploaded = gd_input_audio_file_created_datetime.strftime('%Y-%m-%d-%H%M%S%f')

            processed_files_count += 1
            st.write(f"Starting file {processed_files_count}.")
            st.write(f"Filename: {gd_input_audio_file_name}")

            # Download the original file to local repo (before any conversion)
            input_audio_local_path = gd_download_file(gd_input_audio_file_id, gd_input_audio_file_name)
            st.write(f"Downloaded file: {gd_input_audio_file_name} with MIME type: {gd_input_audio_file_mimeType}")

            # Convert the input file to MP3 with the same name. Delete the input file
            output_mp3_local_path, seconds_transcribed = convert_to_mp3(input_audio_local_path, gd_input_audio_file_mimeType)
            if output_mp3_local_path:
                st.write(f"Converted {gd_input_audio_file_name} to .mp3 format for transcription. Output file: {output_mp3_local_path}. Seconds converted: {seconds_transcribed}.")
            else:
                st.write("Conversion failed. Unsupported MIME type or an error occurred.")

            # Generate new file name based on timestamp and rename file
            datetime_transcribed = datetime.now().strftime('%Y-%m-%d-%H%M%S%f') 
            gd_output_mp3_file_name = f"SIGNAL_{datetime_transcribed}.mp3"
            renamed_mp3_local_path = rename_file(output_mp3_local_path, gd_output_mp3_file_name)
            st.write(f"Renamed {output_mp3_local_path} to {gd_output_mp3_file_name}")

            # Upload mp3 file to Google Drive
            gd_output_mp3_file_id = gd_upload_file(renamed_mp3_local_path, GD_FOLDER_ID_TRANSCRIBED_AUDIO,mime_type='audio/mpeg')
            st.write(f".mp3 file uploaded to Google Drive with ID: {gd_output_mp3_file_id}")

            # Transcribe the audio
            raw_transcription = transcribe(renamed_mp3_local_path)
            st.write(f"Raw transcription generated for {gd_output_mp3_file_name}.")

            # Prompt GPT-4 to format the transcription
            system_prompt = (
                "Optimize the raw audio transcription by formatting and cleaning up the text for readability."
                "Preserve all details to ensure no information is lost."
                "Focus on clarity and direct communication of key points."
                "For any parts of the text that you think is a name of a company or person, mark them as (*name?)."
            )
            task_prompt = raw_transcription
            processed_transcription = openai_llm_call(system_prompt, task_prompt, client)
           
            # After extracting the formatted transcription content from GPT-4
            st.write(f"Formatted transcription generated for {gd_output_mp3_file_name}")

            # Create the docx
            gd_output_mp3_file_link = gd_get_shareable_link(gd_output_mp3_file_id)
            gd_transcript_file_name = os.path.join(os.getcwd(), f"SIGNAL_{datetime_transcribed}_TRANSCRIPT_UNTAGGED.docx")

            try:
                doc = Document()
                doc.add_heading('NOS - Daily Digest Transcription', 0)
                doc.add_paragraph("*Please note that names will be difficult to identify in this iteration of development.\nThe next phase of development will target entity resolution. If there is any confusion, please reference the source audio file and alter the transcribed text.")
                doc.add_paragraph("*Please rename the file with the initials of the recorder and confirmation you have reviewed.")
                doc.add_heading('Transcribed on:', level=1)
                doc.add_paragraph(f"{datetime_transcribed}")
                doc.add_heading('Recorded on:', level=1)
                doc.add_paragraph(f"{datetime_uploaded}")
                doc.add_heading('Seconds transcribed:')
                doc.add_paragraph(str(seconds_transcribed))
                doc.add_heading('MP3 File Link:')
                doc.add_paragraph(gd_output_mp3_file_link) 
                doc.add_heading('Raw Transcription:', level=1)
                doc.add_paragraph(raw_transcription)
                doc.add_heading('Formatted Transcription:', level=1)
                doc.add_paragraph(processed_transcription)
                
                # Save the document
                doc.save(gd_transcript_file_name)
                st.write(f"Generated .docx Transcript.")
            except Exception as e:
                st.write(f"Error creating document: {str(e)}")

            # Upload the docx
            if os.path.exists(gd_transcript_file_name):
                gd_transcript_file_id = gd_upload_file(
                    gd_transcript_file_name,
                    GD_FOLDER_ID_TRANSCRIBED_TEXT,
                    mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                st.write(f"Transcript .docx uploaded to Google Drive with ID: {gd_transcript_file_id}")

                # Update the file's properties directly
                properties = {
                    'transcription_timestamp': datetime_transcribed,
                    'upload_timestamp': datetime_uploaded,
                    'duration_seconds': str(seconds_transcribed),
                    'raw_audio_file_link': gd_input_audio_file_link,
                    'mp3_file_link': gd_output_mp3_file_link
                }
                gd_update_file_properties(gd_transcript_file_id, properties)
                st.write(f"Updated properties for file ID: {gd_transcript_file_id}. Properties are {properties}")
            else:
                st.write(f"Document not found at {gd_transcript_file_name}. Skipping upload.")

            # Move the original audio file from the GDrive to archive folder
            gd_move_file_between_folders(gd_input_audio_file_id, GD_FOLDER_ID_PROCESSED_RAW_AUDIO)
            st.write(f"Moved {gd_input_audio_file_name} to archive folder.")
            
            # Clean up the local files after all processing
            #TODO have everything in a temp dir and clear it.
            #  
            # Delete the original upload file if it exists
            if os.path.exists(input_audio_local_path):
                os.remove(input_audio_local_path)
                st.write(f"Deleted original upload file: {input_audio_local_path}")

            # Delete the converted .mp3 file if it exists
            if os.path.exists(renamed_mp3_local_path):
                os.remove(renamed_mp3_local_path)
                st.write(f"Deleted local .mp3 file: {renamed_mp3_local_path}")

            if os.path.exists(gd_transcript_file_name):
                os.remove(gd_transcript_file_name)
                st.write(f"Deleted local .docx file: {gd_transcript_file_name}")

            gd_transcript_file_link = gd_get_shareable_link(gd_transcript_file_id)
            st.write(f"File {processed_files_count} complete. Transcript Link: {gd_transcript_file_link}")

            # Write to ingress log

            # Clean the transcription texts to remove line breaks and extra whitespace
            raw_transcription_single_line = re.sub(r'\s+', ' ', raw_transcription).strip()
            processed_transcription_single_line = re.sub(r'\s+', ' ', processed_transcription).strip()

            # Prepare the row data with cleaned transcription texts
            row = [
                gd_transcript_file_id,
                datetime_transcribed,
                datetime_uploaded,
                seconds_transcribed,
                gd_transcript_file_link,
                os.path.basename(gd_transcript_file_name),
                gd_output_mp3_file_id,
                gd_output_mp3_file_name,
                gd_output_mp3_file_link,
                gd_input_audio_file_id,
                gd_input_audio_file_name,
                gd_input_audio_file_link,
                gd_input_audio_file_mimeType,
                raw_transcription_single_line,
                processed_transcription_single_line
            ]

            # Append the row to the spreadsheet
            try:
                request = sheets_service.spreadsheets().values().append(
                    spreadsheetId=GD_SPREADSHEET_ID_INGRESS_LOG,
                    range=f'{GD_SHEET_NAME_INGRESS_LOG}!A:Z',
                    valueInputOption='RAW',
                    insertDataOption='INSERT_ROWS',
                    body={'values': [row]}
                )
                response = request.execute()
                st.write(f"Logged transcription to spreadsheet.")
            except Exception as e:
                st.write(f"Error writing to spreadsheet: {str(e)}")

    except Exception as e:
        st.error(f"Error during transcription: {str(e)}")

    st.success(f"{processed_files_count} transcription(s) complete! Find files in the folder linked below.")
    st.markdown('[Transcriptions Folder](https://drive.google.com/drive/u/0/folders/1HVT-YrVNnMy4ag0h6hqawl2PVef-Fc0C)')
