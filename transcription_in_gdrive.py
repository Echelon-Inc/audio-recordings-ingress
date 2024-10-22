#Open AI API Key is pulling from terminal command: export OPENAI_API_KEY = 'apikey'
import streamlit as st
from google.oauth2 import service_account
from googleapiclient.discovery import build
from openai import OpenAI
from dotenv import load_dotenv
from datetime import datetime
from docx import Document
from pydub import AudioSegment
import os
import io
import json
import re
import ffmpeg #had to brew install
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload

# Load environment variables from the .env file
# load_dotenv()
# OpenAI.api_key = os.getenv("OPENAI_API_KEY")
OpenAI.api_key = st.secrets["OPENAI_API_KEY"]

# Define scopes and load credentials
SCOPES = ['https://www.googleapis.com/auth/drive', 'https://www.googleapis.com/auth/documents']
creds = service_account.Credentials.from_service_account_file('credentials.json', scopes=SCOPES)

# Initialize the Drive API client
drive_service = build('drive', 'v3', credentials=creds)
docs_service = build('docs', 'v1', credentials=creds)
client = OpenAI()

# Define folder IDs
unprocessed_audio_gd_folder_id = '10asUMD9jFbWlIXsTxqSezPdJkJU8czdm'
transcribed_audio_gd_folder_id = '1KfdDf2LR7abUn-TpG9MrjYv3fhGXHmox'
transcribed_text_gd_folder_id = '1HVT-YrVNnMy4ag0h6hqawl2PVef-Fc0C'
trash_gd_folder_id = '1TZzr1cxQGxohvFRR63kip7PxMCWExwTR'

def convert_to_mp3(input_file, output_file):
    audio = AudioSegment.from_file(input_file)
    audio.export(output_file, format="mp3")

# Function to list .mp3 or .mpg files in a folder
def list_audio_files(folder_id):
    query = f"'{folder_id}' in parents and (mimeType contains 'audio/' or mimeType contains 'video/')"
    results = drive_service.files().list(q=query, fields="files(id, name, mimeType)").execute()
    files = results.get('files', [])
    # Debugging: Print the file types found
    for file in files:
        print(f"Found file: {file['name']} with MIME type: {file['mimeType']}")
    return files

# Function to move the file to another folder (trash folder)
def move_file_to_trash(file_id, trash_folder_id):
    try:
        # Retrieve the existing parents (folders) of the file
        file = drive_service.files().get(fileId=file_id, fields='parents').execute()
        previous_parents = ",".join(file.get('parents'))

        # Debugging: Print file ID and parent folder info
        print(f"Moving file {file_id} from parents {previous_parents} to trash folder {trash_folder_id}")

        # Attempt to move the file to the trash folder
        updated_file = drive_service.files().update(
            fileId=file_id,
            addParents=trash_folder_id,
            removeParents=previous_parents,
            fields='id, parents'
        ).execute()

        print(f"File ID {file_id} moved to trash folder ID {trash_folder_id}")
    except Exception as e:
        print(f"Error moving file {file_id}: {str(e)}")
        raise  # Re-raise the exception for further debugging

# Function to convert .mpg to .mp3 using ffmpeg
def convert_mpg_to_mp3(mpg_file_path, mp3_file_path):
    try:
        # Debugging: Print conversion details
        print(f"Converting {mpg_file_path} to {mp3_file_path}")
        ffmpeg.input(mpg_file_path).output(mp3_file_path).run()
        print(f"Converted {mpg_file_path} to {mp3_file_path}")
        return mp3_file_path
    except ffmpeg.Error as e:
        print(f"Error converting file: {e}")
        return None

# Function to download a file from Google Drive
def download_file(file_id, file_name):
    request = drive_service.files().get_media(fileId=file_id)
    fh = io.FileIO(file_name, 'wb')
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while done is False:
        status, done = downloader.next_chunk()
        print(f"Download {file_name}: {int(status.progress() * 100)}%.")
    return file_name

# Function to upload a file to Google Drive
def upload_file_to_drive(file_path, folder_id, mime_type='audio/mpeg'):
    file_metadata = {
        'name': os.path.basename(file_path),
        'parents': [folder_id]
    }
    media = MediaFileUpload(file_path, mimetype=mime_type)
    uploaded_file = drive_service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id'
    ).execute()
    print(f"Uploaded file {file_path} to Google Drive with ID: {uploaded_file.get('id')}")
    return uploaded_file.get('id')

# Transcribe audio using OpenAI Whisper
def transcribe(audio_file_path):
    with open(audio_file_path, 'rb') as audio_file:
        transcription = client.audio.transcriptions.create(
            model="whisper-1", 
            file=audio_file,
        )
    return transcription.text

def convert_m4a_to_mp3(m4a_file_path, mp3_file_path):
    try:
        print(f"Converting {m4a_file_path} to {mp3_file_path}")
        ffmpeg.input(m4a_file_path).output(mp3_file_path).run()
        print(f"Converted {m4a_file_path} to {mp3_file_path}")
        return mp3_file_path
    except ffmpeg.Error as e:
        print(f"Error converting .m4a file: {str(e)}")
        return None

# Move file to a different Google Drive folder
def move_file_to_folder(file_id, target_folder_id):
    try:
        # Retrieve the existing parents to remove
        file = drive_service.files().get(fileId=file_id, fields='parents').execute()
        previous_parents = ",".join(file.get('parents'))

        # Move the file to the new folder
        updated_file = drive_service.files().update(
            fileId=file_id,
            addParents=target_folder_id,
            removeParents=previous_parents,
            fields='id, parents'
        ).execute()

        print(f"File ID {file_id} moved to folder ID {target_folder_id}")
    except Exception as e:
        print(f"Error moving file {file_id}: {str(e)}")

# Function to create a shareable link for a file
def get_shareable_link(file_id):
    try:
        # Update file permissions to make it shareable (optional: adjust to 'anyoneWithLink' for broader access)
        permission = {
            'type': 'anyone',
            'role': 'reader'
        }
        drive_service.permissions().create(fileId=file_id, body=permission).execute()
        
        # Get the shareable link
        file = drive_service.files().get(fileId=file_id, fields='webViewLink').execute()
        return file.get('webViewLink')
    except Exception as e:
        print(f"Error getting shareable link for file {file_id}: {str(e)}")
        return None

#Save Medatada as JSON
def save_metadata_as_json(metadata, json_file_path):
    """
    Save the given metadata dictionary as a .json file.
    """
    try:
        with open(json_file_path, 'w') as json_file:
            json.dump(metadata, json_file, indent=4)
        print(f"Metadata saved to {json_file_path}")
    except Exception as e:
        print(f"Error saving metadata to JSON: {str(e)}")

#Upload file to Gdrive
def upload_file_to_drive(file_path, folder_id, mime_type):
    file_metadata = {
        'name': os.path.basename(file_path),
        'parents': [folder_id]
    }
    media = MediaFileUpload(file_path, mimetype=mime_type)
    uploaded_file = drive_service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id'
    ).execute()
    print(f"Uploaded file {file_path} to Google Drive with ID: {uploaded_file.get('id')}")
    return uploaded_file.get('id')

# Streamlit UI
st.set_page_config(
        page_title="NOS Speech2Text",
        page_icon="Echelon_Icon_Sky Blue.png",
)
st.image("Echelon_Icon_Sky Blue.png", caption="The Home for Aliens", width = 125)
st.title("NOS Daily Digest Transcription App - Custom Built for Kerri Faber")
st.write("Once you have uploaded your files to the folder linked below, click the 'Transcribe Audio Files' button to transcribe. Full instructions are available on Notion.")
st.markdown('[Upload Folder](https://drive.google.com/drive/folders/10asUMD9jFbWlIXsTxqSezPdJkJU8czdm?usp=drive_link)')
st.markdown('[Notion](https://www.notion.so/Pulse-4799295f90594380b55f75e0d78dbb03?p=11b9668a26d680e39d57e8243d8f7178&pm=s)')

# Add a reset button
if st.button('Reset App'):
    st.query_params.clear()  # Simulate a reset by clearing query parameters

if st.button('Transcribe Audio Files'):
    st.write("Transcription started...")

    try:
        # Step 1: List all .mp3 and .mpg files in the unprocessed audio folder
        audio_files = list_audio_files(unprocessed_audio_gd_folder_id)
        count = 0
        file_count = len(audio_files)
        st.write(f"Found {file_count} audio files to transcribe.")

        for file in audio_files:
            input_audio_gd_file_id = file['id']
            input_audio_file_name = file['name']  # Original file name
            input_audio_mime_type = file['mimeType']
            count += 1
            st.write(f"Starting file {count}.")

            # Step 2: Download the original file (before any conversion)
            input_audio_path = download_file(input_audio_gd_file_id, input_audio_file_name)
            st.write(f"Downloaded file: {input_audio_file_name} with MIME type: {input_audio_mime_type}")

            # Generate new file name based on timestamp
            timestamp = datetime.now().strftime("%Y-%m-%d-%H%M%S")
            final_audio_file_name = f"SIGNAL_{timestamp}.mp3"
            final_audio_path = os.path.join(os.path.dirname(input_audio_path), final_audio_file_name)

            # Step 3: Check if it's an .mpg file and convert it to .mp3 with timestamped name
            print(f"MIME Type: {input_audio_mime_type} ({type(input_audio_mime_type)})")
            if 'video' in str(input_audio_mime_type):
                final_audio_path = convert_mpg_to_mp3(input_audio_path, final_audio_file_name)
                st.write(f"Converted {input_audio_file_name} to .mp3 format for transcription. Renamed to {final_audio_file_name}")
            else: 
                # Step 3: If it's already an .mp3, rename the file. Path is the same
                convert_to_mp3(input_audio_path, final_audio_path)
                #os.rename(input_audio_path, final_audio_path)
                st.write(f"Renamed {input_audio_file_name} to {final_audio_file_name}")

            # Now final_audio_path points to the renamed file (either converted or not)

            # Step 4: Transcribe the audio
            raw_transcription = transcribe(final_audio_path)
            st.write(f"Raw transcription for {final_audio_file_name}: {raw_transcription}")

            # Step 5: Upload the .mp3 file (converted or original) to Google Drive
            mp3_gd_file_id = upload_file_to_drive(final_audio_path, transcribed_audio_gd_folder_id,mime_type='audio/mpeg')
            st.write(f".mp3 file uploaded to Google Drive with ID: {mp3_gd_file_id}")

            # Step 6: Move the original file to trash folder
            move_file_to_trash(input_audio_gd_file_id, trash_gd_folder_id)
            st.write(f"Moved {input_audio_file_name} to trash folder.")

            # Step 7: Prompt GPT-4 to format the transcription
            prompt = "Optimize this raw transcription by formatting and cleaning up the text for a reader, while preserving all details. Your perspective should be that of a diligent third party analyzing the transcript presented by your boss, not simply a first-person reformat. It is important you communicate the important components of their message directly."
            print(f"Raw Transcription Type: {type(raw_transcription)}")
            print(f"Prompt to GPT-4: {prompt}")
           
            try:
                # Send the raw transcription to GPT-4 for formatting
                completion = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": prompt},
                        {"role": "user", "content": raw_transcription}
                    ]
                )
                formatted_transcription = completion.choices[0].message.content if completion.choices[0].message else ""
                print(f"Formatted transcription: {formatted_transcription}")
            except Exception as e:
                # Handle any exceptions that occur during the GPT-4 call
                print(f"Error during GPT-4 API call: {str(e)}")
                formatted_transcription = "LLM Processing Failed. Use ChatGPT manually"  # Set 'formatted_text' to None or an empty string to avoid undefined variable error
            
           # After extracting the formatted transcription content from GPT-4
            print(f"ft Type: {type(formatted_transcription)}")

            # Step 8: Prepare the docx
            mp3_link = get_shareable_link(mp3_gd_file_id)

            # Replace invalid characters in the timestamp for the file name
            timestamp = datetime.now().strftime('%Y-%m-%d-%H:%M:%S')  # Safe format for file names

            # Prepare the .docx file name
            doc_file_name = os.path.join(os.getcwd(), f"SIGNAL_{timestamp}_INITIALS_TRANSCRIPTION_FOR REVIEW.docx")

            try:
                doc = Document()
                doc.add_heading('NOS - Daily Digest Transcription', 0)
                doc.add_paragraph("*Please note that names will be difficult to identify in this iteration of development.\nThe next phase of development will target entity resolution. If there is any confusion, please reference the source audio file and alter the transcribed text.")
                doc.add_paragraph("*Please rename the file with the initials of the recorder and confirmation you have reviewed.")
                doc.add_heading('Transcribed on:', level=1)
                doc.add_paragraph(f"{timestamp}")
                doc.add_heading('Recorded by: (please specify)', level=1)
                doc.add_paragraph("David McColl/Erik Allen/Kerri Faber/Joel Moxley/Christian Bader")
                doc.add_heading('Raw Transcription:', level=1)
                doc.add_paragraph(raw_transcription)
                doc.add_heading('Formatted Transcription:', level=1)
                doc.add_paragraph(formatted_transcription)
                doc.add_heading('MP3 File Link:')
                doc.add_paragraph(mp3_link)
                
                # Save the document
                doc.save(doc_file_name)
                print(f"Generated .docx file: {doc_file_name}")
            except Exception as e:
                print(f"Error creating document: {str(e)}")

            # Check if the file exists before attempting to upload
            if os.path.exists(doc_file_name):
                print(f"Document saved successfully at {doc_file_name}")
                
                # Upload the .docx file to Google Drive
                doc_id = upload_file_to_drive(doc_file_name, transcribed_text_gd_folder_id, mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            else:
                print(f"Document not found at {doc_file_name}. Skipping upload.")
            #Step 10: Save metadata in a JSON file
            # metadata = {
            #     "input_audio_file_name": input_audio_file_name,
            #     "input_audio_file_id": input_audio_gd_file_id,
            #     "mp3_file_name": final_audio_file_name,
            #     "mp3_file_id": mp3_gd_file_id,
            #     "transcription_text_raw": raw_transcription,
            #     "transcription_text_formatted": formatted_transcription,
            #     "time_of_transcription": timestamp,
            #     "google_doc_id": doc_id,
            #     "google_doc_link": f"https://docs.google.com/document/d/{doc_id}"
            # }
            # json_file_path = f"{final_audio_file_name.replace('.mp3', '')}_metadata.json"
            # save_metadata_as_json(metadata, json_file_path)
            
            # Step 11: Clean up the local files after all processing
            # Delete the original .mpg file if it exists
            if os.path.exists(input_audio_path):
                os.remove(input_audio_path)
                st.write(f"Deleted original .mpg file: {input_audio_path}")

            # Delete the converted .mp3 file after uploading and transcription
            if os.path.exists(final_audio_path):
                os.remove(final_audio_path)
                st.write(f"Deleted local .mp3 file: {final_audio_path}")

            if os.path.exists(doc_file_name):
                os.remove(doc_file_name)
                st.write(f"Deleted local .docx file: {doc_file_name}")
            
            st.write(f"File {count} complete.")
    except Exception as e:
        st.error(f"Error during transcription: {str(e)}")

    st.success(f"{count} transcription(s) complete! Find files in the folder linked below.")
    st.markdown('[Transcriptions Folder](https://drive.google.com/drive/u/0/folders/1HVT-YrVNnMy4ag0h6hqawl2PVef-Fc0C)')
